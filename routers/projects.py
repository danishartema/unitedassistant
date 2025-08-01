"""
Project management router for creating and managing projects.
"""
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, status
import logging

from models import User, Project, ProjectMember, ProjectRole, GPTModeSession
from schemas import (
    ProjectCreate, ProjectUpdate, ProjectResponse, ProjectCreateResponse, ProjectInvite, 
    APIResponse
)
from dependencies import get_current_active_user
import uuid
from pydantic import BaseModel
# Removed unused imports - no longer needed in projects.py
from fastapi.responses import FileResponse
import tempfile
from docx import Document
import json
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
from database import get_async_db
from datetime import datetime, timezone
# Removed chatbot service import - no longer needed in projects.py

logger = logging.getLogger(__name__)

router = APIRouter()

# Removed duplicate chatbot functionality - these are now handled in assistant.py

# --- Project Export Functions ---

@router.post("/", response_model=ProjectCreateResponse, status_code=201)
async def create_project(
    project_data: ProjectCreate,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_async_db)
):
    """Create a new project (async)."""
    try:
        # Create project ID first
        project_id = str(uuid.uuid4())
        current_time = datetime.now(timezone.utc)
        
        # Use raw SQL to avoid any ORM relationship issues
        from sqlalchemy import text
        await db.execute(
            text("""
            INSERT INTO projects (id, title, description, owner_id, is_active, created_at, updated_at)
            VALUES (:id, :title, :description, :owner_id, :is_active, :created_at, :updated_at)
            """),
            {
                "id": project_id,
                "title": project_data.title,
                "description": project_data.description,
                "owner_id": current_user.id,
                "is_active": True,
                "created_at": current_time,
                "updated_at": current_time
            }
        )
        
        await db.commit()
        
        # Create the 14 phases for this project
        await create_project_phases(db, project_id)
        
        # Create response data without accessing any ORM objects
        response_data = {
            "id": project_id,
            "title": project_data.title,
            "description": project_data.description,
            "owner": {
                "id": current_user.id,
                "email": current_user.email,
                "full_name": current_user.full_name,
                "role": current_user.role,
                "is_active": current_user.is_active,
                "created_at": current_user.created_at
            },
            "is_active": True,
            "created_at": current_time,
            "updated_at": current_time
        }
        
        logger.info(f"Project created: {project_data.title} by {current_user.email}")
        return response_data
        
    except Exception as e:
        logger.error(f"Error creating project: {e}")
        try:
            await db.rollback()
        except:
            pass  # Ignore rollback errors
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to create project"
        )

async def create_project_phases(db: AsyncSession, project_id: str):
    """Create the 14 phases for a new project."""
    try:
        from models import Phase, PhaseStatus
        from datetime import datetime, timezone
        
        current_time = datetime.now(timezone.utc)
        
        # Phase titles and descriptions
        phase_data = [
            (1, "Project Overview", "Create a comprehensive project overview"),
            (2, "Target Audience Analysis", "Define and analyze target audience"),
            (3, "Problem Statement", "Identify and articulate the problem"),
            (4, "Solution Framework", "Develop solution framework"),
            (5, "Value Proposition", "Define unique value proposition"),
            (6, "Market Research", "Conduct market research and analysis"),
            (7, "Competitive Analysis", "Analyze competitors and market positioning"),
            (8, "Business Model", "Define business model and revenue streams"),
            (9, "Marketing Strategy", "Develop comprehensive marketing strategy"),
            (10, "Implementation Plan", "Create detailed implementation plan"),
            (11, "Risk Assessment", "Identify and assess potential risks"),
            (12, "Success Metrics", "Define key performance indicators"),
            (13, "Timeline and Milestones", "Create project timeline and milestones"),
            (14, "Final Review and Recommendations", "Final review and recommendations")
        ]
        
        # Create phases using raw SQL to avoid ORM issues
        from sqlalchemy import text
        for phase_number, title, description in phase_data:
            phase_id = str(uuid.uuid4())
            await db.execute(
                text("""
                    INSERT INTO phases (id, project_id, phase_number, title, description, status, created_at, updated_at)
                    VALUES (:id, :project_id, :phase_number, :title, :description, :status, :created_at, :updated_at)
                """),
                {
                    "id": phase_id,
                    "project_id": project_id,
                    "phase_number": phase_number,
                    "title": title,
                    "description": description,
                    "status": PhaseStatus.NOT_STARTED.value,
                    "created_at": current_time,
                    "updated_at": current_time
                }
            )
        
        await db.commit()
        logger.info(f"Created 14 phases for project {project_id}")
        
    except Exception as e:
        logger.error(f"Error creating phases for project {project_id}: {e}")
        await db.rollback()
        raise

@router.get("/", response_model=list[ProjectResponse])
async def get_user_projects(
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_async_db)
):
    """Get all projects for the current user (async)."""
    try:
        result = await db.execute(
            select(Project)
            .where(Project.owner_id == current_user.id)
        )
        projects = result.scalars().all()
        
        # Convert to response format manually to avoid greenlet issues
        response_projects = []
        for project in projects:
            project_data = {
                "id": project.id,
                "title": project.title,
                "description": project.description,
                "owner": {
                    "id": current_user.id,
                    "email": current_user.email,
                    "full_name": current_user.full_name,
                    "role": current_user.role,
                    "is_active": current_user.is_active,
                    "created_at": current_user.created_at
                },
                "members": [],  # Empty list for now
                "is_active": project.is_active,
                "created_at": project.created_at,
                "updated_at": project.updated_at
            }
            response_projects.append(project_data)
        
        return response_projects
    except Exception as e:
        logger.error(f"Error getting user projects: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve projects"
        )

@router.get("/{project_id}", response_model=ProjectResponse)
async def get_project(
    project_id: str,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_async_db)
):
    """Get a specific project (async)."""
    try:
        project_result = await db.execute(
            select(Project).where(Project.id == project_id)
        )
        project = project_result.scalar_one_or_none()
        
        if not project:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Project not found"
            )
        
        # Check if user is owner
        if project.owner_id != current_user.id:
            # Check if user is a member
            member_result = await db.execute(
                select(ProjectMember).where(
                    ProjectMember.project_id == project_id,
                    ProjectMember.user_id == current_user.id
                )
            )
            member = member_result.scalar_one_or_none()
            if not member:
                raise HTTPException(
                    status_code=status.HTTP_403_FORBIDDEN,
                    detail="Access denied to this project"
                )
        
        # Format response manually
        response_data = {
            "id": project.id,
            "title": project.title,
            "description": project.description,
            "owner": {
                "id": current_user.id,
                "email": current_user.email,
                "full_name": current_user.full_name,
                "role": current_user.role,
                "is_active": current_user.is_active,
                "created_at": current_user.created_at
            },
            "members": [],  # Empty list for now
            "is_active": project.is_active,
            "created_at": project.created_at,
            "updated_at": project.updated_at
        }
        
        return response_data
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error getting project: {e}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Failed to retrieve project"
        )

@router.put("/{project_id}", response_model=ProjectResponse)
async def update_project(
    project_id: str,
    project_data: ProjectUpdate,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_async_db)
):
    project_result = await db.execute(select(Project).where(Project.id == project_id))
    project = project_result.scalar()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    if project.owner_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only project owner can update project")
    if project_data.title is not None:
        project.title = project_data.title
    if project_data.description is not None:
        project.description = project_data.description
    await db.commit()
    await db.refresh(project)
    
    # Format response manually
    response_data = {
        "id": project.id,
        "title": project.title,
        "description": project.description,
        "owner": {
            "id": current_user.id,
            "email": current_user.email,
            "full_name": current_user.full_name,
            "role": current_user.role,
            "is_active": current_user.is_active,
            "created_at": current_user.created_at
        },
        "members": [],  # Empty list for now
        "is_active": project.is_active,
        "created_at": project.created_at,
        "updated_at": project.updated_at
    }
    
    return response_data

@router.delete("/{project_id}", response_model=APIResponse)
async def delete_project(
    project_id: str,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_async_db)
):
    project = await db.execute(select(Project).where(Project.id == project_id))
    project = project.scalar()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    if project.owner_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only project owner can delete project")
    project.is_active = False
    await db.commit()
    logger.info(f"Project deleted: {project.title} by {current_user.email}")
    return APIResponse(success=True, message="Project deleted successfully")

@router.post("/{project_id}/invite", response_model=APIResponse)
async def invite_user_to_project(
    project_id: str,
    invite_data: ProjectInvite,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_async_db)
):
    project = await db.execute(select(Project).where(Project.id == project_id))
    project = project.scalar()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    if project.owner_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only project owner can invite users")
    invited_user = await db.execute(select(User).where(User.email == invite_data.email))
    invited_user = invited_user.scalar()
    if not invited_user:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="User not found")
    existing_member = await db.execute(select(ProjectMember).where(
        ProjectMember.project_id == project_id,
        ProjectMember.user_id == invited_user.id
    ))
    existing_member = existing_member.scalar()
    if existing_member:
        raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="User is already a member of this project")
    db_member = ProjectMember(
        project_id=project_id,
        user_id=invited_user.id,
        role=invite_data.role
    )
    db.add(db_member)
    await db.commit()
    logger.info(f"User {invited_user.email} invited to project {project.title}")
    return APIResponse(success=True, message=f"User {invited_user.email} invited successfully")

@router.delete("/{project_id}/members/{user_id}", response_model=APIResponse)
async def remove_project_member(
    project_id: str,
    user_id: str,
    current_user: User = Depends(get_current_active_user),
    db: AsyncSession = Depends(get_async_db)
):
    project = await db.execute(select(Project).where(Project.id == project_id))
    project = project.scalar()
    if not project:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Project not found")
    if project.owner_id != current_user.id:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Only project owner can remove members")
    member = await db.execute(select(ProjectMember).where(
        ProjectMember.project_id == project_id,
        ProjectMember.user_id == user_id
    ))
    member = member.scalar()
    if not member:
        raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Member not found")
    await db.delete(member)
    await db.commit()
    logger.info(f"Member {user_id} removed from project {project.title}")
    return APIResponse(success=True, message="Member removed successfully")

# --- GPT Mode Session Management ---

# Removed duplicate chatbot functionality - these are now handled in assistant.py

# Removed duplicate chatbot endpoints - these are now handled in assistant.py

@router.get("/{project_id}/export/pdf")
async def export_project_pdf(
    project_id: str,
    db: AsyncSession = Depends(get_async_db),
    user: User = Depends(get_current_active_user)
):
    # Gather all mode summaries for the project
    sessions = await db.execute(select(GPTModeSession).where(GPTModeSession.project_id == project_id))
    sessions = sessions.scalars().all()
    doc = Document()
    doc.add_heading(f"Project {project_id} Summary", 0)
    for session in sessions:
        doc.add_heading(session.mode_name, level=1)
        for qnum, answer in session.answers.items():
            doc.add_paragraph(f"Q{qnum}: {answer}")
        if session.checkpoint_json.get('output_template'):
            doc.add_paragraph("Output Template:")
            doc.add_paragraph(session.checkpoint_json['output_template'])
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
        doc.save(tmp.name)
        tmp_path = tmp.name
    return FileResponse(tmp_path, filename=f"project_{project_id}_summary.docx")

@router.get("/{project_id}/export/json")
async def export_project_json(
    project_id: str,
    db: AsyncSession = Depends(get_async_db),
    user: User = Depends(get_current_active_user)
):
    sessions = await db.execute(select(GPTModeSession).where(GPTModeSession.project_id == project_id))
    sessions = sessions.scalars().all()
    all_data = {}
    for session in sessions:
        all_data[session.mode_name] = {
            "answers": session.answers,
            "checkpoint_json": session.checkpoint_json
        }
    with tempfile.NamedTemporaryFile(delete=False, suffix='.json', mode='w', encoding='utf-8') as tmp:
        json.dump(all_data, tmp, ensure_ascii=False, indent=2)
        tmp_path = tmp.name
    return FileResponse(tmp_path, filename=f"project_{project_id}_summary.json")
