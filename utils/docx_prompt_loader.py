from docx import Document
from typing import List, Tuple
import os

def extract_questions_from_docx(docx_path: str) -> List[str]:
    """Extracts questions from a .docx file. Assumes each question is a separate paragraph or numbered list item."""
    doc = Document(docx_path)
    questions = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and (text.endswith('?') or text.startswith('Q')):
            questions.append(text)
    # Also check for numbered/bulleted lists
    for para in doc.paragraphs:
        if para.style.name.startswith('List') and para.text.strip():
            questions.append(para.text.strip())
    # Remove duplicates
    questions = list(dict.fromkeys(questions))
    return questions

def extract_instructions_from_docx(docx_path: str) -> str:
    """Extracts the main instruction text from a .docx file (all non-question paragraphs)."""
    doc = Document(docx_path)
    instructions = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and not text.endswith('?'):
            instructions.append(text)
    return '\n'.join(instructions)

def find_docx_file_in_folder(folder_path: str) -> str:
    """Finds the first .docx file in a folder and returns its path."""
    for fname in os.listdir(folder_path):
        if fname.lower().endswith('.docx'):
            return os.path.join(folder_path, fname)
    raise FileNotFoundError(f"No .docx file found in {folder_path}") 