"""
Export service for generating documents in various formats.
"""
import os
import json
from datetime import datetime
from typing import Dict, Any
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches
import logging

from config import settings

logger = logging.getLogger(__name__)


class ExportService:
    """Service for exporting projects to various formats."""
    
    @staticmethod
    def ensure_upload_directory():
        """Ensure upload directory exists."""
        os.makedirs(settings.upload_dir, exist_ok=True)
    
    @staticmethod
    async def export_to_pdf(project_data: Dict[str, Any], file_path: str) -> str:
        """Export project to PDF format."""
        try:
            ExportService.ensure_upload_directory()
            
            # Create PDF document
            doc = SimpleDocTemplate(file_path, pagesize=letter)
            story = []
            styles = getSampleStyleSheet()
            
            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            phase_title_style = ParagraphStyle(
                'PhaseTitle',
                parent=styles['Heading2'],
                fontSize=16,
                spaceBefore=20,
                spaceAfter=10
            )
            
            # Add title
            story.append(Paragraph(project_data['title'], title_style))
            story.append(Spacer(1, 20))
            
            # Add description
            if project_data.get('description'):
                story.append(Paragraph(f"<b>Description:</b> {project_data['description']}", styles['Normal']))
                story.append(Spacer(1, 20))
            
            # Add metadata
            story.append(Paragraph(f"<b>Created:</b> {project_data['created_at']}", styles['Normal']))
            story.append(Paragraph(f"<b>Owner:</b> {project_data['owner']['full_name']}", styles['Normal']))
            story.append(Spacer(1, 30))
            
            # Add phases
            for phase in project_data['phases']:
                # Phase title
                story.append(Paragraph(f"Phase {phase['phase_number']}: {phase['title']}", phase_title_style))
                
                # Phase description
                if phase.get('description'):
                    story.append(Paragraph(f"<i>{phase['description']}</i>", styles['Italic']))
                    story.append(Spacer(1, 10))
                
                # User input
                if phase.get('user_input'):
                    story.append(Paragraph("<b>Input:</b>", styles['Normal']))
                    story.append(Paragraph(phase['user_input'], styles['Normal']))
                    story.append(Spacer(1, 10))
                
                # AI response
                if phase.get('ai_response'):
                    story.append(Paragraph("<b>Response:</b>", styles['Normal']))
                    story.append(Paragraph(phase['ai_response'], styles['Normal']))
                else:
                    story.append(Paragraph("<i>Phase not completed</i>", styles['Italic']))
                
                story.append(Spacer(1, 20))
            
            # Build PDF
            doc.build(story)
            
            logger.info(f"PDF export completed: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"PDF export error: {e}")
            raise Exception(f"Failed to export to PDF: {str(e)}")
    
    @staticmethod
    async def export_to_word(project_data: Dict[str, Any], file_path: str) -> str:
        """Export project to Word format."""
        try:
            ExportService.ensure_upload_directory()
            
            # Create Word document
            doc = Document()
            
            # Title
            title = doc.add_heading(project_data['title'], 0)
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Description
            if project_data.get('description'):
                doc.add_paragraph(f"Description: {project_data['description']}")
            
            # Metadata
            doc.add_paragraph(f"Created: {project_data['created_at']}")
            doc.add_paragraph(f"Owner: {project_data['owner']['full_name']}")
            doc.add_paragraph("")  # Empty line
            
            # Phases
            for phase in project_data['phases']:
                # Phase heading
                doc.add_heading(f"Phase {phase['phase_number']}: {phase['title']}", 1)
                
                # Phase description
                if phase.get('description'):
                    p = doc.add_paragraph()
                    p.add_run(phase['description']).italic = True
                
                # User input
                if phase.get('user_input'):
                    doc.add_paragraph("Input:", style='Heading 2')
                    doc.add_paragraph(phase['user_input'])
                
                # AI response
                if phase.get('ai_response'):
                    doc.add_paragraph("Response:", style='Heading 2')
                    doc.add_paragraph(phase['ai_response'])
                else:
                    p = doc.add_paragraph()
                    p.add_run("Phase not completed").italic = True
                
                doc.add_paragraph("")  # Empty line
            
            # Save document
            doc.save(file_path)
            
            logger.info(f"Word export completed: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"Word export error: {e}")
            raise Exception(f"Failed to export to Word: {str(e)}")
    
    @staticmethod
    async def export_to_json(project_data: Dict[str, Any], file_path: str) -> str:
        """Export project to JSON format."""
        try:
            ExportService.ensure_upload_directory()
            
            # Prepare export data
            export_data = {
                "export_info": {
                    "exported_at": datetime.now().isoformat(),
                    "format": "json",
                    "version": "1.0"
                },
                "project": project_data
            }
            
            # Write JSON file
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(export_data, f, indent=2, ensure_ascii=False, default=str)
            
            logger.info(f"JSON export completed: {file_path}")
            return file_path
            
        except Exception as e:
            logger.error(f"JSON export error: {e}")
            raise Exception(f"Failed to export to JSON: {str(e)}")
    
    @staticmethod
    def get_export_file_path(export_id: str, format: str) -> str:
        """Generate file path for export."""
        ExportService.ensure_upload_directory()
        filename = f"export_{export_id}.{format}"
        return os.path.join(settings.upload_dir, filename)
